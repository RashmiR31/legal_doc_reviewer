import os, shutil, uuid
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader, TextLoader
from docx import Document as DocxDocument
import streamlit as st

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except:
    OCR_AVAILABLE = False

def save_uploaded_files(uploaded_files, upload_dir) -> list[str]:
    os.makedirs(upload_dir, exist_ok=True)
    saved_paths = []
    for f in uploaded_files:
        dst = os.path.join(upload_dir, os.path.basename(f.name))
        with open(dst, "wb") as out:
            shutil.copyfileobj(f, out)
        saved_paths.append(dst)
    return saved_paths

def ocr_pdf_to_docs(pdf_path: str, dpi=300):
    if not OCR_AVAILABLE:
        raise RuntimeError("OCR dependencies not installed.")
    pages = convert_from_path(pdf_path, dpi=dpi)
    return [Document(page_content=pytesseract.image_to_string(p), metadata={"source": pdf_path, "page": i+1})
            for i,p in enumerate(pages)]

def load_file_to_docs(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load_and_split()
            if sum(len(d.page_content.strip()) for d in docs) < 40 and OCR_AVAILABLE:
                st.info("Running OCR...")
                docs = ocr_pdf_to_docs(file_path)
        except:
            docs = ocr_pdf_to_docs(file_path) if OCR_AVAILABLE else []
    elif ext in [".docx", ".doc"]:
        try:
            loader = UnstructuredWordDocumentLoader(file_path)
            docs = loader.load()
        except:
            docx = DocxDocument(file_path)
            docs = [Document(page_content="\n".join([p.text for p in docx.paragraphs]), metadata={"source": file_path})]
    elif ext in [".txt", ".md"]:
        loader = TextLoader(file_path, encoding="utf8")
        docs = loader.load()
    else:
        raise ValueError(f"Unsupported type: {ext}")
    for d in docs: d.metadata.setdefault("source", file_path)
    return docs
